from docx import Document
from docx.shared import Inches,Pt
from docx2pdf import convert
import glob,os,sys
import numpy as np
import pandas as pd

class Make_docx:
    def __init__(self,a,cname):
        self.cn = cname
        self.df = pd.DataFrame(a)
        self.df.columns = [self.cn]
        self.doc = Document()

    def make_simple_report(self,fname,fPath):
        des = self.df.describe().reset_index()
        self.doc.add_heading("Describe Table")
        table = self.doc.add_table(rows=(des.shape[0]),cols=des.shape[1])
        table.style = "Table Grid"
        for i,column in enumerate(des):
            for row in range(des.shape[0]):
                table.cell(row,i).text = str(des[column][row])
        img_list = glob.glob("../out_image/%s.*.jpg"%fname)
        for i in img_list:

            if "line" in i:
                self.doc.add_picture(i,width=Inches(6.0))
            else:
                print(i)
                self.doc.add_picture(i,width=Inches(2.5))
        self.doc.save("%s/%s.docx"%(fPath,fname))
if __name__ == "__main__":
    np.random.seed(1)
    aa = np.random.normal(0,2.0,1000)
    p = Make_docx(aa,"test")
    p.make_simple_report("test","../out_document")
